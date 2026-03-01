import logging
from typing import Dict, List
from pathlib import Path
from datetime import datetime
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.colors import HexColor
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

logger = logging.getLogger(__name__)


def wrap_text(text: str, canvas_obj, font_name: str, font_size: int, max_width: int) -> List[str]:
    words = text.split()
    lines = []
    current_line = []
    for word in words:
        test_line = ' '.join(current_line + [word])
        if canvas_obj.stringWidth(test_line, font_name, font_size) <= max_width:
            current_line.append(word)
        else:
            if current_line:
                lines.append(' '.join(current_line))
            current_line = [word]
    if current_line:
        lines.append(' '.join(current_line))
    return lines


def create_packing_slip_pdf(order_data: Dict, config: Dict, barcode_path: str = "", base_path: Path = None) -> str:
    """CONFIG-DRIVEN PDF - V17.0 - courier_note yellow highlight added"""
    if base_path is None:
        base_path = Path.cwd()

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    invoice_no = order_data.get('invoice_no', 'UNKNOWN').replace('# ', '').strip()
    pdf_path = base_path / 'packing_slips' / f"packing_slip_{invoice_no}_{timestamp}.pdf"
    c = canvas.Canvas(str(pdf_path), pagesize=A4)
    width, height = A4
    pdf_config = config.get('pdf', {})
    sections = pdf_config.get('sections', {})
    global_style = pdf_config.get('global_style', {})

    page_bg_rgb = global_style.get('page_bg_color', [255, 255, 255])
    page_bg = HexColor(f'#{page_bg_rgb[0]:02x}{page_bg_rgb[1]:02x}{page_bg_rgb[2]:02x}')
    c.setFillColor(page_bg)
    c.rect(0, 0, width, height, fill=1, stroke=0)

    font_family = global_style.get('font_family', 'Helvetica')
    font_size = global_style.get('font_size', 10)
    font_color_rgb = global_style.get('font_color', [44, 62, 80])
    font_color = HexColor(f'#{font_color_rgb[0]:02x}{font_color_rgb[1]:02x}{font_color_rgb[2]:02x}')
    highlight_color_rgb = global_style.get('highlight_color', [255, 204, 0])
    highlight_color = HexColor(f'#{highlight_color_rgb[0]:02x}{highlight_color_rgb[1]:02x}{highlight_color_rgb[2]:02x}')
    y_position = height - 30

    # ==================== BARCODE ====================
    barcode_config = sections.get('barcode', {})
    barcode_height = barcode_config.get('height', 60)
    barcode_width = barcode_config.get('width', 180)
    barcode_x = barcode_config.get('x', 30)
    barcode_y_position = y_position - barcode_height

    if barcode_config.get('enabled', True) and barcode_path and Path(barcode_path).exists():
        try:
            c.drawImage(barcode_path, barcode_x, barcode_y_position,
                        width=barcode_width, height=barcode_height,
                        preserveAspectRatio=True)
        except Exception as e:
            logger.error(f"[✗] Barcode failed: {e}")

    # ==================== LOGO ====================
    contact_config = sections.get('contact_card', {})
    if contact_config.get('enabled', True):
        assets_dir = base_path / 'assets'
        filename = contact_config.get('filename', 'logo1')
        extensions = contact_config.get('extensions', ['jpeg', 'jpg', 'png'])
        img_path = None
        for ext in extensions:
            for path in [assets_dir / f"{filename}.{ext}", base_path / f"{filename}.{ext}"]:
                if path.exists():
                    img_path = path
                    break
            if img_path:
                break
        if img_path:
            try:
                size_increase = contact_config.get('size_increase_percent', 25) / 100
                new_width = int(contact_config.get('width', 312) * (1 + size_increase))
                new_height = int(contact_config.get('height', 108) * (1 + size_increase))
                logo_x = contact_config.get('x', 240)
                logo_y_position = barcode_y_position - contact_config.get('y_offset', 20)
                c.drawImage(str(img_path), logo_x, logo_y_position,
                            width=new_width, height=new_height,
                            preserveAspectRatio=True)
                logger.info(f"✅ Logo drawn: width={new_width}, height={new_height}, x={logo_x}, y={logo_y_position}")
            except Exception as e:
                logger.error(f"[✗] Logo failed: {e}")

    max_header_height = max(barcode_height, contact_config.get('height', 108) * (1 + contact_config.get('size_increase_percent', 25) / 100)) if contact_config.get('enabled', True) else barcode_height
    y_position -= (max_header_height + contact_config.get('spacing_below', 5))

    # ==================== TITLE ====================
    title_config = sections.get('title', {})
    if title_config.get('enabled', True):
        delivery_date = order_data.get('delivery_date', 'N/A')
        full_title = title_config.get('text', 'FF PICKUP - {delivery_date} | Invoice #: {invoice_no}')
        full_title = full_title.replace('{delivery_date}', delivery_date).replace('{invoice_no}', invoice_no)

        title_font = title_config.get('font_family', 'Helvetica-Bold')
        title_font_size = title_config.get('font_size', 11)
        title_height = title_config.get('height', 25)

        bg_color_rgb = title_config.get('bg_color')
        if bg_color_rgb:
            bg_color = HexColor(f'#{bg_color_rgb[0]:02x}{bg_color_rgb[1]:02x}{bg_color_rgb[2]:02x}')
            c.setFillColor(bg_color)
            c.roundRect(30, y_position - title_height, 540, title_height, 5, fill=1, stroke=0)

        text_color_rgb = title_config.get('text_color', [0, 0, 0])
        c.setFillColor(HexColor(f'#{text_color_rgb[0]:02x}{text_color_rgb[1]:02x}{text_color_rgb[2]:02x}'))
        c.setFont(title_font, title_font_size)

        if title_config.get('alignment', 'center') == 'center':
            c.drawString((width - c.stringWidth(full_title, title_font, title_font_size)) / 2, y_position - 15, full_title)
        else:
            c.drawString(40, y_position - 15, full_title)

        spacing_below = title_config.get('spacing_below', 10)
        y_position -= (title_height + spacing_below + 14)

    # ==================== CUSTOMER INFO ====================
    customer_config = sections.get('customer_info', {})
    if customer_config.get('enabled', True):
        start_x = customer_config.get('start_x', 40)
        value_x_offset = customer_config.get('value_x_offset', 80)
        line_height = customer_config.get('line_height', 14)
        max_value_width = 520 - value_x_offset
        c.setFillColor(font_color)

        is_gift_order = bool(order_data.get('gift_recipient', '').strip())
        gift_recipient_name = order_data.get('gift_recipient', '').strip()

        fields_list = customer_config.get('fields_gift_order', []) if is_gift_order else customer_config.get('fields_normal_order', [])
        line_count = 0

        for field_info in fields_list:
            if not field_info.get('show', True):
                continue

            field_key = field_info.get('key')
            value = order_data.get(field_key, '')

            if 'combined_format' in field_info:
                try:
                    value = field_info['combined_format'].format(**order_data)
                except:
                    pass

            if not value or value == "N/A":
                continue

            label = field_info.get('label', '')
            if field_info.get('label_dynamic', False) and '{gift_recipient}' in label:
                label = label.replace('{gift_recipient}', gift_recipient_name if gift_recipient_name else 'Recipient')

            # ✅ courier_note check
            is_courier_note = (field_key == 'courier_note')

            if label:
                c.setFont(f"{font_family}-Bold", font_size)
                c.setFillColor(font_color)
                c.drawString(start_x, y_position - (line_count * line_height), label)

            if value:
                wrapped_lines = wrap_text(str(value), c, font_family, font_size, max_value_width)
                for wrapped_line in wrapped_lines:
                    y_offset = line_count * line_height
                    if y_position - y_offset < 150:
                        c.showPage()
                        c.setFillColor(page_bg)
                        c.rect(0, 0, width, height, fill=1, stroke=0)
                        y_position = height - 30
                        line_count = 0
                        y_offset = 0
                        c.setFillColor(font_color)

                    # ✅ courier_note ke liye yellow highlight background
                    if is_courier_note:
                        text_width = c.stringWidth(wrapped_line, font_family, font_size)
                        highlight_x = start_x + value_x_offset - 2
                        # Baseline se 3px neeche start, height = font_size + 2 (proper fit)
                        highlight_y = y_position - y_offset - 3
                        highlight_h = font_size + 2
                        c.setFillColor(highlight_color)
                        c.rect(highlight_x, highlight_y, text_width + 4, highlight_h, fill=1, stroke=0)

                    c.setFont(font_family, font_size)
                    c.setFillColor(font_color)
                    c.drawString(start_x + value_x_offset, y_position - y_offset, wrapped_line)
                    line_count += 1

        y_position -= (line_count * line_height + 10)

    # ==================== RECOMMENDATION ====================
    rec_config = sections.get('recommendation', {})
    if rec_config.get('enabled', True):
        is_gift_order = bool(order_data.get('gift_recipient', '').strip())
        start_x = rec_config.get('start_x', 40)
        line_height = rec_config.get('line_height', 14)
        rec_font_size = rec_config.get('font_size', 10)
        max_width = rec_config.get('max_width', 520)
        labels = rec_config.get('labels', {})

        text_lines = []

        if is_gift_order:
            gift_note = order_data.get('gift_note', '').strip()
            if gift_note:
                gift_recipient = order_data.get('gift_recipient', '').strip()
                gift_recipient_name = gift_recipient if gift_recipient else order_data.get('customer_name', '').split()[0]
                message_label = labels.get('message', 'Message')
                text_lines.append(('message_combined', f"{message_label} for {gift_recipient_name}:", gift_note))
        else:
            customer_name = order_data.get('customer_name', 'Customer')
            first_name = customer_name.split()[0] if customer_name else 'Customer'
            greeting = labels.get('greeting', 'Hi {first_name},').replace('{first_name}', first_name)
            text_lines.append(('normal', greeting))

            for line in rec_config.get('generic_message_normal', '').split('\n'):
                if line.strip():
                    text_lines.append(('generic_multiline', line.strip()))

            feedback_response = order_data.get('ai_feedback_response', '').strip()
            show_feedback = rec_config.get('show_feedback_response', 'always')
            if feedback_response and show_feedback in ('normal_order_only', 'always'):
                text_lines.append(('empty', ''))
                text_lines.append(('feedback', feedback_response))

            text_lines.append(('empty', ''))
            ai_recommendation = order_data.get('ai_recommendation', '').strip()
            if ai_recommendation:
                text_lines.append(('rec_combined', labels.get('recommendation', 'PS:'), ai_recommendation))

        c.setFillColor(font_color)
        current_line = 0
        last_generic_line_text = ""

        for line_data in text_lines:
            if y_position - (current_line * line_height) < 150:
                c.showPage()
                c.setFillColor(page_bg)
                c.rect(0, 0, width, height, fill=1, stroke=0)
                y_position = height - 30
                current_line = 0
                c.setFillColor(font_color)

            if len(line_data) == 2:
                line_type, line_text = line_data

                if line_type == 'empty':
                    current_line += 1

                elif line_type == 'feedback':
                    for wrapped_line in wrap_text(line_text, c, f"{font_family}-Bold", rec_font_size, max_width):
                        if y_position - (current_line * line_height) < 150:
                            c.showPage(); c.setFillColor(page_bg); c.rect(0, 0, width, height, fill=1, stroke=0)
                            y_position = height - 30; current_line = 0; c.setFillColor(font_color)
                        c.setFillColor(font_color)
                        c.setFont(f"{font_family}-Bold", rec_font_size)
                        c.drawString(start_x, y_position - (current_line * line_height), wrapped_line)
                        current_line += 1

                elif line_type == 'generic_multiline':
                    for wrapped_line in wrap_text(line_text, c, font_family, rec_font_size, max_width):
                        if y_position - (current_line * line_height) < 150:
                            c.showPage(); c.setFillColor(page_bg); c.rect(0, 0, width, height, fill=1, stroke=0)
                            y_position = height - 30; current_line = 0; c.setFillColor(font_color)
                        c.setFont(font_family, rec_font_size)
                        c.drawString(start_x, y_position - (current_line * line_height), wrapped_line)
                        last_generic_line_text = wrapped_line
                        current_line += 1

                    if line_text.strip().startswith("Nathan"):
                        heart_config = rec_config.get('heart_image', {})
                        if heart_config.get('enabled', False):
                            assets_dir = base_path / 'assets'
                            filename = heart_config.get('filename', 'heart')
                            extensions = heart_config.get('extensions', ['jpg', 'jpeg', 'png'])
                            img_path = None
                            for ext in extensions:
                                for path in [assets_dir / f"{filename}.{ext}", base_path / f"{filename}.{ext}"]:
                                    if path.exists():
                                        img_path = path
                                        break
                                if img_path:
                                    break
                            if img_path:
                                try:
                                    text_width = c.stringWidth(last_generic_line_text, font_family, rec_font_size)
                                    c.drawImage(str(img_path),
                                                start_x + text_width + 3,
                                                y_position - ((current_line - 1) * line_height) - 2,
                                                width=heart_config.get('width', 12),
                                                height=heart_config.get('height', 12),
                                                preserveAspectRatio=True, mask='auto')
                                    logger.info(f"✅ Heart image added")
                                except Exception as e:
                                    logger.warning(f"⚠️ Heart image failed: {e}")
                else:
                    for wrapped_line in wrap_text(line_text, c, font_family, rec_font_size, max_width):
                        if y_position - (current_line * line_height) < 150:
                            c.showPage(); c.setFillColor(page_bg); c.rect(0, 0, width, height, fill=1, stroke=0)
                            y_position = height - 30; current_line = 0; c.setFillColor(font_color)
                        c.setFont(font_family, rec_font_size)
                        c.drawString(start_x, y_position - (current_line * line_height), wrapped_line)
                        current_line += 1

            elif len(line_data) == 3:
                line_type, label_text, value_text = line_data

                if line_type == 'message_combined':
                    c.setFont(f"{font_family}-Bold", rec_font_size)
                    c.drawString(start_x, y_position - (current_line * line_height), label_text)
                    label_width = c.stringWidth(label_text, f"{font_family}-Bold", rec_font_size)

                    for idx, wrapped_line in enumerate(wrap_text(value_text, c, font_family, rec_font_size, max_width - label_width - 10)):
                        if y_position - (current_line * line_height) < 150:
                            c.showPage(); c.setFillColor(page_bg); c.rect(0, 0, width, height, fill=1, stroke=0)
                            y_position = height - 30; current_line = 0; c.setFillColor(font_color)
                        c.setFillColor(font_color)
                        c.setFont(font_family, rec_font_size)
                        x_pos = start_x + label_width + 5 if idx == 0 else start_x
                        c.drawString(x_pos, y_position - (current_line * line_height), wrapped_line)
                        current_line += 1

                elif line_type == 'rec_combined':
                    c.setFont(f"{font_family}-Bold", rec_font_size)
                    c.drawString(start_x, y_position - (current_line * line_height), label_text)
                    current_line += 1

                    for wrapped_line in wrap_text(value_text, c, font_family, rec_font_size, max_width):
                        if y_position - (current_line * line_height) < 150:
                            c.showPage(); c.setFillColor(page_bg); c.rect(0, 0, width, height, fill=1, stroke=0)
                            y_position = height - 30; current_line = 0; c.setFillColor(font_color)
                        c.setFillColor(font_color)
                        c.setFont(font_family, rec_font_size)
                        c.drawString(start_x, y_position - (current_line * line_height), wrapped_line)
                        current_line += 1

        y_position -= (current_line * line_height + 10)

    # ==================== TABLE & RECYCLE LOGO ====================
    table_recycle_config = sections.get('table_and_recycle', {})
    if table_recycle_config.get('enabled', True):
        table_config = table_recycle_config.get('table', {})
        recycle_config = table_recycle_config.get('recycle_logo', {})
        table_start_y = y_position

        table_x = table_config.get('x', 30)
        table_width = table_config.get('width', 270)
        row_height = table_config.get('row_height', 20)
        header_config = table_config.get('header', {})
        rows_config = table_config.get('rows', {})
        columns_config = table_config.get('columns', [])

        header_bg_rgb = header_config.get('bg_color', [44, 62, 80])
        header_bg = HexColor(f'#{header_bg_rgb[0]:02x}{header_bg_rgb[1]:02x}{header_bg_rgb[2]:02x}')
        header_text_rgb = header_config.get('text_color', [255, 255, 255])
        header_text_color = HexColor(f'#{header_text_rgb[0]:02x}{header_text_rgb[1]:02x}{header_text_rgb[2]:02x}')
        header_font = header_config.get('font_family', 'Helvetica-Bold')
        header_font_size = header_config.get('font_size', 10)
        header_height = header_config.get('height', 25)

        if recycle_config.get('enabled', True):
            assets_dir = base_path / 'assets'
            filename = recycle_config.get('filename', 'recycle1')
            extensions = recycle_config.get('extensions', ['jpeg', 'jpg', 'png'])
            img_path = None
            for ext in extensions:
                for path in [assets_dir / f"{filename}.{ext}", base_path / f"{filename}.{ext}"]:
                    if path.exists():
                        img_path = path
                        break
                if img_path:
                    break
            if img_path:
                try:
                    recycle_x = recycle_config.get('x', 320)
                    size_reduction = recycle_config.get('size_reduction_percent', 15) / 100
                    recycle_width = int(recycle_config.get('width', 250) * (1 - size_reduction))
                    recycle_height = int(recycle_config.get('height', 250) * (1 - size_reduction))
                    c.drawImage(str(img_path), recycle_x, table_start_y - recycle_height,
                                width=recycle_width, height=recycle_height, preserveAspectRatio=True)
                except Exception as e:
                    logger.error(f"[✗] Recycle logo failed: {e}")

        c.setFillColor(header_bg)
        c.roundRect(table_x, y_position - header_height, table_width, header_height, 5, fill=1, stroke=0)
        c.setFillColor(header_text_color)
        c.setFont(header_font, header_font_size)
        x_offset = 0
        for col in columns_config:
            c.drawString(table_x + x_offset + 5, y_position - 15, col.get('label', ''))
            x_offset += col.get('width', 50)
        y_position -= (header_height + 5)

        even_bg = HexColor(f'#{rows_config.get("even_bg", [248,249,250])[0]:02x}{rows_config.get("even_bg", [248,249,250])[1]:02x}{rows_config.get("even_bg", [248,249,250])[2]:02x}')
        odd_bg = HexColor(f'#{rows_config.get("odd_bg", [255,255,255])[0]:02x}{rows_config.get("odd_bg", [255,255,255])[1]:02x}{rows_config.get("odd_bg", [255,255,255])[2]:02x}')
        row_font = rows_config.get('font_family', 'Helvetica')
        row_font_size = rows_config.get('font_size', 9)
        line_spacing = rows_config.get('line_spacing', 14)

        for idx, item in enumerate(order_data.get('items', [])):
            if y_position < 100:
                c.showPage()
                c.setFillColor(page_bg)
                c.rect(0, 0, width, height, fill=1, stroke=0)
                y_position = height - 30
                c.setFillColor(header_bg)
                c.roundRect(table_x, y_position - header_height, table_width, header_height, 5, fill=1, stroke=0)
                c.setFillColor(header_text_color)
                c.setFont(header_font, header_font_size)
                x_offset = 0
                for col in columns_config:
                    c.drawString(table_x + x_offset + 5, y_position - 15, col.get('label', ''))
                    x_offset += col.get('width', 50)
                y_position -= (header_height + 5)

            c.setFillColor(even_bg if idx % 2 == 0 else odd_bg)
            c.roundRect(table_x, y_position - (row_height - 2), table_width, row_height - 2, 3, fill=1, stroke=0)
            c.setFillColor(font_color)
            x_offset = 0
            for col in columns_config:
                col_key = col.get('key', '')
                value = str(item.get(col_key, ''))[:col.get('max_chars', 50)]
                c.setFont(f"{row_font}-Bold" if col_key == "qty" else row_font, row_font_size)
                c.drawString(table_x + x_offset + 5, y_position - (line_spacing - 1), value)
                x_offset += col.get('width', 50)
            y_position -= row_height

        y_position -= table_config.get('spacing_below', 0)

    # ==================== FOOTER ====================
    footer_config = sections.get('footer', {})
    if footer_config.get('enabled', True):
        footer_height = footer_config.get('height', 30)
        if y_position < footer_height + 60:
            c.showPage()
            c.setFillColor(page_bg)
            c.rect(0, 0, width, height, fill=1, stroke=0)
            y_position = height - 30

        y_position -= 2
        total_qty = sum(int(item.get('qty', 0)) for item in order_data.get('items', []))

        footer_bg_rgb = footer_config.get('bg_color')
        if footer_bg_rgb:
            c.setFillColor(HexColor(f'#{footer_bg_rgb[0]:02x}{footer_bg_rgb[1]:02x}{footer_bg_rgb[2]:02x}'))
            c.roundRect(30, y_position - footer_height, 270, footer_height, 5, fill=1, stroke=0)

        footer_text_rgb = footer_config.get('text_color', [0, 0, 0])
        c.setFillColor(HexColor(f'#{footer_text_rgb[0]:02x}{footer_text_rgb[1]:02x}{footer_text_rgb[2]:02x}'))
        footer_font = footer_config.get('font_family', 'Helvetica-Bold')
        footer_font_size = footer_config.get('font_size', 11)
        c.setFont(footer_font, footer_font_size)

        footer_template = footer_config.get('template', {})
        if isinstance(footer_template, dict):
            left_text = footer_template.get('left', '').replace('{total_qty}', str(total_qty))
            right_text = footer_template.get('right', '')
            c.drawString(40, y_position - (footer_height // 2) + 3, left_text)
            right_width = c.stringWidth(right_text, footer_font, footer_font_size)
            c.drawString(300 - right_width - 10, y_position - (footer_height // 2) + 3, right_text)
        else:
            c.drawString(45, y_position - (footer_height // 2) + 3,
                         footer_template.replace('{total_qty}', str(total_qty)).replace('{invoice_no}', invoice_no))

        y_position -= (footer_height + 15)

    # ==================== CONTACT LINE ====================
    contact_line_config = sections.get('contact_line', {})
    if contact_line_config.get('enabled', False):
        contact_text = contact_line_config.get('text', '')
        contact_font = contact_line_config.get('font_family', 'Helvetica')
        contact_font_size = contact_line_config.get('font_size', 9)
        contact_color_rgb = contact_line_config.get('text_color', [0, 0, 0])
        c.setFillColor(HexColor(f'#{contact_color_rgb[0]:02x}{contact_color_rgb[1]:02x}{contact_color_rgb[2]:02x}'))
        c.setFont(contact_font, contact_font_size)
        margin_bottom = contact_line_config.get('margin_bottom', 20)
        if contact_line_config.get('alignment', 'center') == 'center':
            c.drawString((width - c.stringWidth(contact_text, contact_font, contact_font_size)) / 2, margin_bottom, contact_text)
        else:
            c.drawString(40, margin_bottom, contact_text)

    c.save()
    logger.info(f"✅ PDF Created: {pdf_path}")
    return str(pdf_path)


def create_packing_slip_docx(order_data: Dict, config: Dict, barcode_path: str = "", base_path: Path = None) -> str:
    """CONFIG-DRIVEN DOCX - V13.0 - courier_note yellow highlight added"""
    if base_path is None:
        base_path = Path.cwd()

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    invoice_no = order_data.get('invoice_no', 'UNKNOWN').replace('# ', '').strip()
    docx_path = base_path / 'packing_slips' / f"packing_slip_{invoice_no}_{timestamp}.docx"

    try:
        doc = Document()
        for section in doc.sections:
            section.page_height = Inches(11.69)
            section.page_width = Inches(8.27)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)

        pdf_config = config.get('pdf', {})
        sections_config = pdf_config.get('sections', {})
        global_style = pdf_config.get('global_style', {})

        # ==================== BARCODE & LOGO ====================
        barcode_config = sections_config.get('barcode', {})
        contact_config = sections_config.get('contact_card', {})

        if barcode_config.get('enabled', True) or contact_config.get('enabled', True):
            header_table = doc.add_table(rows=1, cols=2)
            header_table.alignment = WD_TABLE_ALIGNMENT.CENTER

            if barcode_config.get('enabled', True) and barcode_path and Path(barcode_path).exists():
                try:
                    barcode_para = header_table.rows[0].cells[0].paragraphs[0]
                    barcode_para.add_run().add_picture(barcode_path, width=Inches(3.6))
                except Exception as e:
                    logger.warning(f"[DOCX] Barcode insert failed: {e}")

            if contact_config.get('enabled', True):
                assets_dir = base_path / 'assets'
                filename = contact_config.get('filename', 'logo1')
                extensions = contact_config.get('extensions', ['jpeg', 'jpg', 'png'])
                img_path = None
                for ext in extensions:
                    for path in [assets_dir / f"{filename}.{ext}", base_path / f"{filename}.{ext}"]:
                        if path.exists():
                            img_path = path
                            break
                    if img_path:
                        break
                if img_path:
                    try:
                        contact_para = header_table.rows[0].cells[1].paragraphs[0]
                        contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                        contact_para.add_run().add_picture(str(img_path), width=Inches(3.6))
                    except Exception as e:
                        logger.warning(f"[DOCX] Logo insert failed: {e}")

        doc.add_paragraph()

        # ==================== TITLE ====================
        title_config = sections_config.get('title', {})
        if title_config.get('enabled', True):
            delivery_date = order_data.get('delivery_date', 'N/A')
            full_title = title_config.get('text', 'FF PICKUP - {delivery_date} | Invoice #: {invoice_no}')
            full_title = full_title.replace('{delivery_date}', delivery_date).replace('{invoice_no}', invoice_no)

            title_para = doc.add_paragraph()
            if title_config.get('alignment', 'center') == 'center':
                title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            title_run = title_para.add_run(full_title)
            title_run.bold = True
            title_run.font.size = Pt(title_config.get('font_size', 11))
            text_color_rgb = title_config.get('text_color', [0, 0, 0])
            title_run.font.color.rgb = RGBColor(*text_color_rgb)

            bg_color_rgb = title_config.get('bg_color')
            if bg_color_rgb:
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), f'{bg_color_rgb[0]:02x}{bg_color_rgb[1]:02x}{bg_color_rgb[2]:02x}')
                title_para._element.get_or_add_pPr().append(shading_elm)

        doc.add_paragraph()
        doc.add_paragraph()

        # ==================== CUSTOMER INFO ====================
        customer_config = sections_config.get('customer_info', {})
        if customer_config.get('enabled', True):
            is_gift_order = bool(order_data.get('gift_recipient', '').strip())
            gift_recipient_name = order_data.get('gift_recipient', '').strip()
            fields_list = customer_config.get('fields_gift_order', []) if is_gift_order else customer_config.get('fields_normal_order', [])

            for field_info in fields_list:
                if not field_info.get('show', True):
                    continue

                field_key = field_info.get('key')
                value = order_data.get(field_key, '')

                if 'combined_format' in field_info:
                    try:
                        value = field_info['combined_format'].format(**order_data)
                    except:
                        pass

                if not value or value == "N/A":
                    continue

                label = field_info.get('label', '')
                if field_info.get('label_dynamic', False) and '{gift_recipient}' in label:
                    label = label.replace('{gift_recipient}', gift_recipient_name if gift_recipient_name else 'Recipient')

                para = doc.add_paragraph()
                if label:
                    label_run = para.add_run(f"{label} ")
                    label_run.bold = True
                    label_run.font.size = Pt(global_style.get('font_size', 10))

                value_run = para.add_run(str(value))
                value_run.font.size = Pt(global_style.get('font_size', 10))

                # ✅ courier_note ke liye yellow highlight DOCX
                if field_key == 'courier_note':
                    rPr = value_run._r.get_or_add_rPr()
                    highlight_elm = OxmlElement('w:highlight')
                    highlight_elm.set(qn('w:val'), 'yellow')
                    rPr.append(highlight_elm)

        doc.add_paragraph()

        # ==================== RECOMMENDATION ====================
        rec_config = sections_config.get('recommendation', {})
        if rec_config.get('enabled', True):
            is_gift_order = bool(order_data.get('gift_recipient', '').strip())
            labels = rec_config.get('labels', {})

            if is_gift_order:
                gift_note = order_data.get('gift_note', '').strip()
                if gift_note:
                    gift_recipient = order_data.get('gift_recipient', '').strip()
                    gift_recipient_name = gift_recipient if gift_recipient else order_data.get('customer_name', '').split()[0]
                    message_para = doc.add_paragraph()
                    label_run = message_para.add_run(f"{labels.get('message', 'Message')} for {gift_recipient_name}: ")
                    label_run.bold = True
                    label_run.font.size = Pt(rec_config.get('font_size', 10))
                    value_run = message_para.add_run(gift_note)
                    value_run.font.size = Pt(rec_config.get('font_size', 10))
            else:
                customer_name = order_data.get('customer_name', 'Customer')
                first_name = customer_name.split()[0] if customer_name else 'Customer'
                greeting = labels.get('greeting', 'Hi {first_name},').replace('{first_name}', first_name)

                greeting_para = doc.add_paragraph(greeting)
                greeting_para.runs[0].font.size = Pt(rec_config.get('font_size', 10))

                last_para = None
                for line in rec_config.get('generic_message_normal', '').split('\n'):
                    if line.strip():
                        msg_para = doc.add_paragraph(line.strip())
                        msg_para.runs[0].font.size = Pt(rec_config.get('font_size', 10))
                        last_para = msg_para

                if last_para:
                    heart_config = rec_config.get('heart_image', {})
                    if heart_config.get('enabled', False):
                        assets_dir = base_path / 'assets'
                        filename = heart_config.get('filename', 'heart')
                        extensions = heart_config.get('extensions', ['jpg', 'jpeg', 'png'])
                        img_path = None
                        for ext in extensions:
                            for path in [assets_dir / f"{filename}.{ext}", base_path / f"{filename}.{ext}"]:
                                if path.exists():
                                    img_path = path
                                    break
                            if img_path:
                                break
                        if img_path:
                            try:
                                last_para.add_run(' ').add_picture(str(img_path), width=Inches(heart_config.get('width', 12) / 72))
                                logger.info("✅ Heart image added to DOCX")
                            except Exception as e:
                                logger.warning(f"⚠️ DOCX heart image failed: {e}")

                feedback_response = order_data.get('ai_feedback_response', '').strip()
                show_feedback = rec_config.get('show_feedback_response', 'always')
                if feedback_response and show_feedback in ('normal_order_only', 'always'):
                    doc.add_paragraph()
                    feedback_para = doc.add_paragraph()
                    feedback_run = feedback_para.add_run(feedback_response)
                    feedback_run.bold = True
                    feedback_run.font.size = Pt(rec_config.get('font_size', 10))

                ai_recommendation = order_data.get('ai_recommendation', '').strip()
                if ai_recommendation:
                    doc.add_paragraph()
                    rec_para = doc.add_paragraph()
                    rec_label_run = rec_para.add_run(f"{labels.get('recommendation', 'PS:')} ")
                    rec_label_run.bold = True
                    rec_label_run.font.size = Pt(rec_config.get('font_size', 10))
                    rec_value = rec_para.add_run(ai_recommendation)
                    rec_value.font.size = Pt(rec_config.get('font_size', 10))

        doc.add_paragraph()

        # ==================== TABLE & RECYCLE LOGO ====================
        table_recycle_config = sections_config.get('table_and_recycle', {})
        if table_recycle_config.get('enabled', True):
            table_config = table_recycle_config.get('table', {})
            recycle_config = table_recycle_config.get('recycle_logo', {})

            container_table = doc.add_table(rows=1, cols=2)
            container_table.alignment = WD_TABLE_ALIGNMENT.LEFT

            products_cell = container_table.rows[0].cells[0]
            columns_config = table_config.get('columns', [])
            items = order_data.get('items', [])

            if items and columns_config:
                products_table = products_cell.add_table(rows=1, cols=len(columns_config))
                products_table.style = 'Light Grid Accent 1'

                header_config = table_config.get('header', {})
                hdr_cells = products_table.rows[0].cells
                for idx, col in enumerate(columns_config):
                    hdr_cells[idx].text = col.get('label', '')
                    hdr_run = hdr_cells[idx].paragraphs[0].runs[0]
                    hdr_run.bold = True
                    hdr_run.font.size = Pt(header_config.get('font_size', 10))
                    text_color_rgb = header_config.get('text_color', [255, 255, 255])
                    hdr_run.font.color.rgb = RGBColor(*text_color_rgb)
                    bg_color_rgb = header_config.get('bg_color', [44, 62, 80])
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), f'{bg_color_rgb[0]:02x}{bg_color_rgb[1]:02x}{bg_color_rgb[2]:02x}')
                    hdr_cells[idx]._element.get_or_add_tcPr().append(shading_elm)

                rows_config = table_config.get('rows', {})
                even_bg_rgb = rows_config.get('even_bg', [248, 249, 250])
                odd_bg_rgb = rows_config.get('odd_bg', [255, 255, 255])

                for idx, item in enumerate(items):
                    row_cells = products_table.add_row().cells
                    bg_rgb = even_bg_rgb if idx % 2 == 0 else odd_bg_rgb
                    for col_idx, col in enumerate(columns_config):
                        col_key = col.get('key', '')
                        value = str(item.get(col_key, ''))[:col.get('max_chars', 50)]
                        row_cells[col_idx].text = value
                        cell_run = row_cells[col_idx].paragraphs[0].runs[0]
                        cell_run.font.size = Pt(rows_config.get('font_size', 9))
                        if col_key == "qty":
                            cell_run.bold = True
                        shading_elm = OxmlElement('w:shd')
                        shading_elm.set(qn('w:fill'), f'{bg_rgb[0]:02x}{bg_rgb[1]:02x}{bg_rgb[2]:02x}')
                        row_cells[col_idx]._element.get_or_add_tcPr().append(shading_elm)

            if recycle_config.get('enabled', True):
                recycle_cell = container_table.rows[0].cells[1]
                recycle_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                assets_dir = base_path / 'assets'
                filename = recycle_config.get('filename', 'recycle1')
                extensions = recycle_config.get('extensions', ['jpeg', 'jpg', 'png'])
                img_path = None
                for ext in extensions:
                    for path in [assets_dir / f"{filename}.{ext}", base_path / f"{filename}.{ext}"]:
                        if path.exists():
                            img_path = path
                            break
                    if img_path:
                        break
                if img_path:
                    try:
                        recycle_para = recycle_cell.paragraphs[0]
                        recycle_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        recycle_para.add_run().add_picture(str(img_path), width=Inches(2.8))
                    except Exception as e:
                        logger.warning(f"[DOCX] Recycle logo insert failed: {e}")

        doc.add_paragraph()

        # ==================== FOOTER ====================
        footer_config = sections_config.get('footer', {})
        if footer_config.get('enabled', True):
            total_qty = sum(int(item.get('qty', 0)) for item in order_data.get('items', []))
            footer_para = doc.add_paragraph()
            footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            footer_template = footer_config.get('template', {})
            if isinstance(footer_template, dict):
                footer_text = f"{footer_template.get('left', '').replace('{total_qty}', str(total_qty))}     {footer_template.get('right', '')}"
            else:
                footer_text = footer_template.replace('{total_qty}', str(total_qty)).replace('{invoice_no}', invoice_no)

            footer_run = footer_para.add_run(footer_text)
            footer_run.bold = True
            footer_run.font.size = Pt(footer_config.get('font_size', 11))
            text_color_rgb = footer_config.get('text_color', [0, 0, 0])
            footer_run.font.color.rgb = RGBColor(*text_color_rgb)

            bg_color_rgb = footer_config.get('bg_color')
            if bg_color_rgb:
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), f'{bg_color_rgb[0]:02x}{bg_color_rgb[1]:02x}{bg_color_rgb[2]:02x}')
                footer_para._element.get_or_add_pPr().append(shading_elm)

        # ==================== CONTACT LINE ====================
        contact_line_config = sections_config.get('contact_line', {})
        if contact_line_config.get('enabled', False):
            contact_para = doc.add_paragraph(contact_line_config.get('text', ''))
            if contact_line_config.get('alignment', 'center') == 'center':
                contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            contact_para.runs[0].font.size = Pt(contact_line_config.get('font_size', 9))

        doc.save(str(docx_path))
        logger.info(f"✅ DOCX Created: {docx_path}")
        return str(docx_path)

    except Exception as e:
        logger.error(f"❌ DOCX Creation failed: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return ""